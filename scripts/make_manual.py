"""Generate RE:Node user manual as .docx"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

BLUE      = RGBColor(0x1d, 0x4e, 0xd8)
DBLUE     = RGBColor(0x1e, 0x40, 0xaf)
MBLUE     = RGBColor(0x37, 0x51, 0xcb)
WHITE     = RGBColor(0xff, 0xff, 0xff)
GRAY      = RGBColor(0x64, 0x74, 0x8b)
LGRAY     = RGBColor(0x94, 0xa3, 0xb8)
FONT      = '맑은 고딕'

def set_font(run, size=10, bold=False, color=None):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color

def h1(text):
    p = doc.add_heading(level=1)
    p.clear()
    run = p.add_run(text)
    set_font(run, 16, True, BLUE)
    return p

def h2(text):
    p = doc.add_heading(level=2)
    p.clear()
    run = p.add_run(text)
    set_font(run, 13, True, DBLUE)
    return p

def h3(text):
    p = doc.add_heading(level=3)
    p.clear()
    run = p.add_run(text)
    set_font(run, 11, True, MBLUE)
    return p

def body(text, bold=False, indent=False, color=None):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text)
    set_font(run, 10, bold, color)
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    run = p.add_run(text)
    set_font(run, 10)
    return p

def sep():
    doc.add_paragraph()

def shade_cell(cell, fill_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def table(headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hr = t.rows[0]
    for i, h in enumerate(headers):
        cell = hr.cells[i]
        shade_cell(cell, '1D4ED8')
        cell.paragraphs[0].clear()
        run = cell.paragraphs[0].add_run(h)
        set_font(run, 9, True, WHITE)
    for ri, row in enumerate(rows):
        fill = 'EFF6FF' if ri % 2 == 0 else 'FFFFFF'
        tr = t.rows[ri + 1]
        for ci, val in enumerate(row):
            cell = tr.cells[ci]
            shade_cell(cell, fill)
            cell.paragraphs[0].clear()
            run = cell.paragraphs[0].add_run(val)
            set_font(run, 9)
    return t

# ─── Cover ───────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('RE:Node 사용자 매뉴얼')
set_font(run, 28, True, BLUE)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run('KoreaIT Academy  AI 기반 지식 그래프 학습 플랫폼')
set_font(run2, 13, False, GRAY)

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = p3.add_run('Version 1.0  |  2026년 4월')
set_font(run3, 10, False, LGRAY)

doc.add_page_break()

# ─── 1. 플랫폼 소개 ──────────────────────────────────────
h1('1. 플랫폼 소개')
body('RE:Node는 KoreaIT Academy의 AI 기반 적응형 학습 플랫폼입니다. 강사가 업로드한 강의 텍스트를 Google Gemini AI가 분석하여 지식 그래프를 자동 생성하고, 학생이 원하는 개념을 선택하면 맞춤형 복습 과제를 즉시 생성합니다.')
sep()

h2('1.1 주요 특징')
bullet('AI 자동 지식 그래프 생성: 강의 텍스트에서 개념 노드와 연결 관계를 자동 추출')
bullet('맞춤형 과제 생성: 단답형, 개념 설명, 비교·대조, 요약, 미니퀴즈 5가지 유형')
bullet('실시간 참여도 분석: 매니저가 학생별 과제 생성·제출 현황을 실시간으로 확인')
bullet('모바일 완전 지원: 스와이프 사이드바, 반응형 레이아웃')
sep()

h2('1.2 사용자 역할')
table(
    ['역할', '설명', '주요 권한'],
    [
        ['학생 (Student)', '강의를 수강하는 학습자', '지식 그래프 탐색, 과제 생성·제출'],
        ['강사 (Instructor)', '강의를 개설하는 교수자', '강의 생성, 텍스트 업로드, 학생 현황 조회'],
        ['매니저 (Manager)', '학습 현황을 관리하는 담당자', '전체 참여도 분석, 학생별 드릴다운'],
        ['관리자 (Admin)', '플랫폼 전체 관리자', '모든 기능 접근'],
    ]
)
sep()

# ─── 2. 시작하기 ─────────────────────────────────────────
doc.add_page_break()
h1('2. 시작하기')

h2('2.1 접속 방법')
body('웹 브라우저(Chrome 권장)에서 아래 URL로 접속합니다.')
body('https://revmap-frontend-1004217999763.asia-northeast3.run.app', bold=True, indent=True)
sep()

h2('2.2 로그인')
bullet('이메일과 비밀번호를 입력하고 [로그인] 버튼을 클릭합니다.')
bullet('로그인 화면 좌측에는 RE:Node 지식 그래프 애니메이션이 실시간 표시됩니다.')
bullet('계정이 없으면 [회원가입] 링크를 클릭해 계정을 생성합니다.')
sep()

h2('2.3 테스트 계정 (모든 계정 초기 비밀번호: 1234)')
table(
    ['역할', '이메일'],
    [
        ['관리자', 'admin@renode.io'],
        ['강사1 (Python/Java/C/C++/SQL/자료구조)', 'teacher1@renode.io'],
        ['강사2 (네트워크/OS/알고리즘/컴퓨터구조)', 'teacher2@renode.io'],
        ['학생1 (teacher1 수강생)', 'student1@renode.io'],
        ['학생2 (teacher1 수강생)', 'student2@renode.io'],
        ['학생3 (teacher2 수강생)', 'student3@renode.io'],
        ['학생4 (teacher2 수강생)', 'student4@renode.io'],
        ['매니저', 'manager1@renode.io'],
    ]
)
sep()

h2('2.4 모바일 사용')
bullet('상단 햄버거 메뉴(≡)를 탭하면 사이드바가 열립니다.')
bullet('사이드바를 왼쪽으로 드래그하거나 바깥 영역을 탭하면 닫힙니다.')

# ─── 3. 학생 매뉴얼 ──────────────────────────────────────
doc.add_page_break()
h1('3. 학생 매뉴얼')

h2('3.1 대시보드')
bullet('로그인 후 내 강의 현황 카드(전체 강의, 활성 그래프)를 확인합니다.')
bullet('[내 강의 →] 버튼으로 수강 강의 목록으로 이동합니다.')
bullet('최근 강의 목록에서 강의 상세 페이지로 바로 접근할 수 있습니다.')
sep()

h2('3.2 강의 탐색')
bullet('[내 강의] 또는 [강의 목록] 메뉴에서 수강 강의를 확인합니다.')
bullet('강의 카드 클릭 → 강의 상세 페이지로 이동합니다.')
sep()

h2('3.3 지식 그래프 탐색')
bullet('각 원(노드)이 하나의 개념을 나타냅니다.')
bullet('노드를 클릭하면 개념 설명과 연결된 개념을 확인할 수 있습니다.')
bullet('복습하고 싶은 개념을 클릭하여 선택합니다 (선택 시 색상 변경).')
bullet('여러 개념을 동시에 선택하면 통합 과제를 생성할 수 있습니다.')
bullet('우측 [전체 개념] 패널의 버튼으로도 개념을 선택할 수 있습니다.')
sep()

h2('3.4 AI 과제 생성')
bullet('개념 노드를 1개 이상 선택합니다.')
bullet('우측 패널에서 [과제 유형]을 선택합니다.')
body('과제 유형: 단답형 / 개념 설명 / 비교·대조 / 요약 / 미니퀴즈', indent=True)
bullet('[난이도]를 선택합니다 (쉬움 / 보통 / 어려움).')
bullet('[과제 생성] 버튼 클릭 → AI가 30초 이내에 과제를 생성합니다.')
sep()

h2('3.5 답안 제출')
bullet('생성된 과제 내용을 확인합니다.')
bullet('텍스트 영역에 답안을 작성합니다.')
bullet('[제출] 버튼을 클릭하면 답안이 저장됩니다.')
bullet('제출된 답안은 담당 강사가 확인할 수 있습니다.')

# ─── 4. 강사 매뉴얼 ──────────────────────────────────────
doc.add_page_break()
h1('4. 강사 매뉴얼')

h2('4.1 강의 생성')
bullet('[내 강의] 메뉴 → [새 강의] 버튼 클릭')
bullet('강의 제목과 설명을 입력하고 [생성] 클릭')
sep()

h2('4.2 강의 텍스트 업로드 및 그래프 생성')
bullet('강의 상세 페이지 상단 [강의 텍스트 업로드] 패널을 펼칩니다.')
bullet('강의 내용 텍스트를 붙여넣기 합니다 (최대 50,000자).')
bullet('[업로드] 버튼 클릭 후 [그래프 생성] 버튼을 클릭합니다.')
bullet('AI가 개념과 연결 관계를 자동으로 추출합니다.')
bullet('생성 완료(상태: 완료) 후 지식 그래프가 화면에 표시됩니다.')
sep()

h2('4.3 학생 과제 현황 조회')
body('강의 상세 페이지 상단 탭에서 [학생 과제 현황]을 클릭합니다.')
bullet('학생 목록: 이름, 과제 생성 수, 제출 수, 참여 상태를 확인합니다.')
bullet('학생 이름 클릭: 해당 학생의 과제 목록이 펼쳐집니다.')
bullet('과제 항목 클릭: 생성된 과제 내용과 제출 답안을 동시에 확인합니다.')
body('제출 여부: 초록 체크(✓) = 제출완료 / 회색 원(○) = 미제출', indent=True)

# ─── 5. 매니저 매뉴얼 ────────────────────────────────────
doc.add_page_break()
h1('5. 매니저 매뉴얼')

h2('5.1 분석 대시보드')
bullet('좌측 메뉴 [분석] 클릭')
bullet('상단 요약 카드: 강사 수 / 전체 학생 수 / 참여 학생 수 / 전체 참여율')
sep()

h2('5.2 강사별 현황')
bullet('강사 카드 클릭 → 해당 강사의 강의 목록이 펼쳐집니다.')
bullet('각 강의의 수강생 수, 참여 학생, 과제 수, 제출 수, 참여율을 확인합니다.')
body('참여율 색상: 70% 이상(초록) / 40~70%(노랑) / 40% 미만(빨강)', indent=True)
sep()

h2('5.3 학생별 참여도 드릴다운')
bullet('강의 행 클릭 → 수강생 참여 현황이 펼쳐집니다.')
body('참여 수준 기준:', indent=True, bold=True)
body('  - 높음: 과제 3건 이상 & 제출 2건 이상', indent=True)
body('  - 보통: 과제·제출 각 1건 이상', indent=True)
body('  - 낮음: 과제 생성만 있음', indent=True)
body('  - 미참여: 과제·제출 모두 없음', indent=True)
bullet('학생별 과제 생성 수, 제출 수, 마지막 활동 시간을 확인합니다.')
bullet('활동 막대 그래프로 상대적 참여도를 시각화합니다.')

# ─── 6. 강의 목록 ────────────────────────────────────────
doc.add_page_break()
h1('6. 현재 등록 강의 목록')

h2('선생1 담당 강의')
table(
    ['강의명', '노드', '주요 내용'],
    [
        ['Python 프로그래밍: 문법과 알고리즘', '35', '기본 문법, OOP, 자료구조, 정렬/탐색, DP'],
        ['Java 프로그래밍: 객체지향과 핵심 개념', '26', 'JVM, OOP, 제네릭, 컬렉션, 람다/스트림'],
        ['C 언어 프로그래밍: 기초부터 시스템까지', '26', '포인터, 메모리, 구조체, 파일 I/O'],
        ['C++ 프로그래밍: 객체지향과 현대 C++', '26', '스마트포인터, STL, 템플릿, 이동시맨틱'],
        ['SQL과 데이터베이스: 기초부터 심화까지', '24', 'DDL/DML, JOIN, 인덱스, 트랜잭션, 정규화'],
        ['자료구조: 핵심 개념과 구현', '23', '배열, 연결리스트, 트리, 힙, 해시, 그래프'],
    ]
)
sep()

h2('선생2 담당 강의')
table(
    ['강의명', '노드', '주요 내용'],
    [
        ['컴퓨터 네트워크: 프로토콜과 인터넷 구조', '23', 'OSI 7계층, TCP/IP, HTTP, DNS, TLS'],
        ['운영체제: 프로세스, 메모리, 파일시스템', '23', '프로세스, 스케줄링, 가상메모리, 동기화'],
        ['알고리즘: 설계와 분석', '25', '정렬, 그래프탐색, 최단경로, DP, 탐욕'],
        ['컴퓨터 구조: CPU, 메모리, 명령어', '20', 'CPU, ALU, 파이프라인, 캐시, 메모리계층'],
    ]
)
sep()

# ─── 7. FAQ ──────────────────────────────────────────────
doc.add_page_break()
h1('7. 자주 묻는 질문 (FAQ)')

h3('Q. 과제 생성이 오래 걸려요.')
body('A. AI 서버 상태에 따라 처음 요청 시 30초까지 걸릴 수 있습니다. 503 오류 시 자동 재시도되므로 페이지를 닫지 말고 기다려 주세요.')
sep()

h3('Q. 지식 그래프가 표시되지 않아요.')
body('A. 강의 상태가 "처리중"이면 그래프 생성 중입니다. 새로고침(↻) 버튼으로 상태를 확인하세요. "대기" 상태라면 강사에게 그래프 생성을 요청하세요.')
sep()

h3('Q. 비밀번호를 변경하고 싶어요.')
body('A. 현재 버전에서는 비밀번호 변경이 관리자를 통해서만 가능합니다. admin@renode.io 로 문의하세요.')
sep()

h3('Q. 매니저는 학생 답안을 볼 수 있나요?')
body('A. 매니저는 참여율·수치 통계만 확인 가능합니다. 과제 내용과 답안은 해당 강사만 조회할 수 있습니다.')
sep()

# ─── 8. 문의 ─────────────────────────────────────────────
h1('8. 문의 및 지원')
bullet('플랫폼 관리자: admin@renode.io')
bullet('GitHub Issues: 프로젝트 저장소 Issues 탭')
sep()
body('본 매뉴얼은 RE:Node v1.0 기준으로 작성되었습니다.', color=LGRAY)

doc.save('docs/ReNode_사용자매뉴얼.docx')
print('Done: docs/ReNode_사용자매뉴얼.docx')
